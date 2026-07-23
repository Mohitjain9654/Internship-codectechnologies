# function load_document(file_path):

#     if file ends with ".pdf":
#         open file with PyPDF2
#         for each page → extract text → append to result
    
#     elif file ends with ".docx":
#         open file with python-docx
#         for each paragraph → get text → append to result
    
#     else:
#         raise error "Unsupported file type"
    
#     return full_text as one string

"""
document_loader.py
==================
STAGE 1 - STEP 1: Document Loading & Text Extraction

What this does:
- Accepts a file path (PDF or DOCX)
- Extracts all text from it
- Returns one clean string

How it works internally:
- PDF: PyPDF2 reads the text layer page by page
- DOCX: python-docx walks through every paragraph
- Both paths return the same thing: a raw text string
"""

import os
import PyPDF2
from docx import Document


def load_document(file_path: str) -> str:
    """
    Main entry point. Detects file type and routes to correct extractor.

    Args:
        file_path: Full path to your .pdf or .docx file

    Returns:
        All extracted text as one string

    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file type isn't PDF or DOCX
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        return _extract_from_pdf(file_path)
    elif extension == ".docx":
        return _extract_from_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file type: '{extension}'. Only .pdf and .docx are supported."
        )


def _extract_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF, page by page.

    Note: Only works on PDFs with a text layer.
    Scanned/image PDFs will return empty — OCR needed for those.
    """
    extracted_pages = []

    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        total_pages = len(reader.pages)
        print(f"[document_loader] PDF has {total_pages} pages")

        for page_num in range(total_pages):
            try:
                page = reader.pages[page_num]
                text = page.extract_text()

                if text and text.strip():
                    extracted_pages.append(text)
                else:
                    print(f"  [Warning] Page {page_num + 1} has no extractable text (may be scanned image)")

            except Exception as e:
                print(f"  [Warning] Skipping page {page_num + 1}: {e}")
                continue

    if not extracted_pages:
        raise ValueError(
            "No text extracted from PDF. It might be a scanned image-based PDF."
        )

    full_text = "\n".join(extracted_pages)
    print(f"[document_loader] Extracted {len(full_text)} characters from PDF")
    return full_text


def _extract_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX, paragraph by paragraph.

    Note: A .docx is actually a ZIP of XML files.
    python-docx parses the XML and gives us paragraph objects.
    """
    doc = Document(file_path)
    extracted_paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            extracted_paragraphs.append(text)

    if not extracted_paragraphs:
        raise ValueError("No text extracted from DOCX file.")

    full_text = "\n".join(extracted_paragraphs)
    print(f"[document_loader] Extracted {len(full_text)} characters from DOCX")
    return full_text


def get_document_info(file_path: str) -> dict:
    """
    Returns metadata about a document without full extraction.
    Used later by the Streamlit UI to show file details.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    file_size_kb = round(os.path.getsize(file_path) / 1024, 2)

    info = {
        "file_name": os.path.basename(file_path),
        "file_type": ext,
        "file_size_kb": file_size_kb,
    }

    if ext == ".pdf":
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            info["page_count"] = len(reader.pages)
    elif ext == ".docx":
        doc = Document(file_path)
        info["paragraph_count"] = len(doc.paragraphs)

    return info


# ── Quick test — run: python document_loader.py your_file.pdf ───────────────
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python document_loader.py <path_to_file>")
        sys.exit(1)

    path = sys.argv[1]

    print("\n=== Document Info ===")
    info = get_document_info(path)
    for k, v in info.items():
        print(f"  {k}: {v}")

    print("\n=== Extracted Text (first 500 chars) ===")
    text = load_document(path)
    print(text[:500])
    print(f"\nTotal characters: {len(text)}")